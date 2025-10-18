import logging
import time
import asyncio
from typing import List, Dict, Any
from pathlib import Path
from io import BytesIO
import fitz  # PyMuPDF for better PDF handling
from docx import Document
from PyPDF2 import PdfReader
import pandas as pd
import json
from concurrent.futures import ThreadPoolExecutor
import psutil
import gc
 
# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
 
class FileProcessingConfig:
    """Configuration for file processing"""
    def __init__(self):
        self.max_file_size_mb = 500  # Maximum file size in MB
        self.chunk_size_mb = 10      # Process files in chunks of this size
        self.memory_threshold = 80   # Stop processing if memory usage exceeds this %
        self.pdf_use_pymupdf = True  # Use PyMuPDF for better text extraction
        self.excel_chunk_size = 1000 # Rows per chunk for Excel files
        self.csv_chunk_size = 5000   # Rows per chunk for CSV files
       
config = FileProcessingConfig()
 
def check_system_resources() -> Dict[str, float]:
    """Check current system resource usage"""
    memory = psutil.virtual_memory()
    cpu = psutil.cpu_percent(interval=0.1)
   
    return {
        "memory_percent": memory.percent,
        "memory_available_gb": memory.available / (1024**3),
        "cpu_percent": cpu
    }
 
def validate_file_size(file_path_or_bytes, filename: str) -> bool:
    """Validate file size before processing"""
    try:
        if hasattr(file_path_or_bytes, 'seek'):
            # File-like object
            file_path_or_bytes.seek(0, 2)  # Seek to end
            size_bytes = file_path_or_bytes.tell()
            file_path_or_bytes.seek(0)  # Reset to beginning
        else:
            # File path
            size_bytes = Path(file_path_or_bytes).stat().st_size
           
        size_mb = size_bytes / (1024 * 1024)
       
        if size_mb > config.max_file_size_mb:
            logging.warning(f"File {filename} ({size_mb:.2f}MB) exceeds maximum size ({config.max_file_size_mb}MB)")
            return False
           
        logging.info(f"File {filename} size: {size_mb:.2f}MB - OK")
        return True
       
    except Exception as e:
        logging.error(f"Error validating file size for {filename}: {e}")
        return False
 
def extract_text(file, filename: str) -> str:
    """
    Enhanced text extraction with large file support and memory management
    """
    logging.info(f"Starting text extraction for file: {filename}")
   
    # Check system resources before processing
    resources = check_system_resources()
    if resources["memory_percent"] > config.memory_threshold:
        logging.warning(f"High memory usage ({resources['memory_percent']:.1f}%) - proceeding with caution")
   
    # Validate file size
    if not validate_file_size(file, filename):
        raise ValueError(f"File {filename} is too large to process")
   
    start_time = time.time()
    ext = filename.split(".")[-1].lower()
    logging.info(f"Detected file extension: {ext}")
 
    try:
        if ext == "pdf":
            result = extract_pdf_text(file, filename)
        elif ext == "docx":
            result = extract_docx_text(file, filename)
        elif ext == "txt":
            result = extract_txt_text(file, filename)
        elif ext == "json":
            result = extract_json_text(file, filename)
        elif ext == "csv":
            result = extract_csv_text(file, filename)
        elif ext in {"xlsx", "xls"}:
            result = extract_excel_text(file, filename)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
 
        end_time = time.time()
        logging.info(f"Text extraction completed for file: {filename} in {end_time - start_time:.2f} seconds")
       
        # Memory cleanup
        gc.collect()
       
        return result
 
    except Exception as e:
        logging.error(f"Error during text extraction for file: {filename} - {e}")
        raise
 
def extract_pdf_text(file, filename: str) -> str:
    """Enhanced PDF text extraction with memory management"""
    logging.info(f"Processing PDF: {filename}")
   
    try:
        if config.pdf_use_pymupdf:
            # Use PyMuPDF for better text extraction and large file handling
            if hasattr(file, 'read'):
                pdf_bytes = file.read()
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            else:
                doc = fitz.open(file)
           
            text_parts = []
            total_pages = len(doc)
            logging.info(f"Processing {total_pages} pages from PDF")
           
            for page_num in range(total_pages):
                try:
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                   
                    if page_text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                   
                    # Memory management - clean up page
                    page = None
                   
                    # Check memory usage periodically
                    if page_num % 50 == 0 and page_num > 0:
                        resources = check_system_resources()
                        if resources["memory_percent"] > config.memory_threshold:
                            logging.warning(f"High memory usage at page {page_num}: {resources['memory_percent']:.1f}%")
                            gc.collect()
                   
                except Exception as e:
                    logging.warning(f"Error processing page {page_num + 1}: {e}")
                    continue
           
            doc.close()
            result = "\n\n".join(text_parts)
           
        else:
            # Fallback to PyPDF2
            reader = PdfReader(file)
            text_parts = []
           
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception as e:
                    logging.warning(f"Error processing page {page_num + 1}: {e}")
                    continue
           
            result = "\n\n".join(text_parts)
       
        logging.info(f"Extracted {len(result)} characters from PDF")
        return result
       
    except Exception as e:
        logging.error(f"Error processing PDF {filename}: {e}")
        raise
 
def extract_excel_text(file, filename: str) -> str:
    """Enhanced Excel processing with chunked reading for large files"""
    logging.info(f"Processing Excel file: {filename}")
   
    try:
        text_parts = []
       
        # Get all sheet names first
        if hasattr(file, 'read'):
            file_bytes = BytesIO(file.read())
            xlsx_file = pd.ExcelFile(file_bytes)
        else:
            xlsx_file = pd.ExcelFile(file)
       
        sheet_names = xlsx_file.sheet_names
        logging.info(f"Found {len(sheet_names)} sheets: {sheet_names}")
       
        for sheet_name in sheet_names:
            logging.info(f"Processing sheet: {sheet_name}")
           
            try:
                # Try to read the entire sheet first to get row count
                sheet_df = pd.read_excel(xlsx_file, sheet_name=sheet_name, nrows=0)
               
                # Read in chunks for large sheets
                chunk_num = 0
                sheet_text_parts = [f"[Sheet: {sheet_name}]"]
               
                # Add headers
                if not sheet_df.columns.empty:
                    headers = " | ".join(str(col) for col in sheet_df.columns)
                    sheet_text_parts.append(f"Headers: {headers}")
               
                # Process data in chunks
                for chunk_df in pd.read_excel(
                    xlsx_file,
                    sheet_name=sheet_name,
                    chunksize=config.excel_chunk_size
                ):
                    chunk_num += 1
                    logging.info(f"Processing chunk {chunk_num} of sheet {sheet_name} ({len(chunk_df)} rows)")
                   
                    # Convert chunk to text
                    chunk_text = dataframe_to_text(chunk_df, f"{sheet_name}_chunk_{chunk_num}")
                    sheet_text_parts.append(chunk_text)
                   
                    # Memory check
                    if chunk_num % 5 == 0:
                        resources = check_system_resources()
                        if resources["memory_percent"] > config.memory_threshold:
                            logging.warning(f"High memory usage processing Excel: {resources['memory_percent']:.1f}%")
                            gc.collect()
               
                text_parts.append("\n".join(sheet_text_parts))
               
            except Exception as e:
                logging.error(f"Error processing sheet {sheet_name}: {e}")
                continue
       
        result = "\n\n".join(text_parts)
        logging.info(f"Extracted {len(result)} characters from Excel file")
        return result
       
    except Exception as e:
        logging.error(f"Error processing Excel file {filename}: {e}")
        raise
 
def extract_csv_text(file, filename: str) -> str:
    """Enhanced CSV processing with chunked reading"""
    logging.info(f"Processing CSV file: {filename}")
   
    try:
        text_parts = [f"[CSV File: {filename}]"]
        chunk_num = 0
       
        # Process CSV in chunks to handle large files
        for chunk_df in pd.read_csv(file, chunksize=config.csv_chunk_size):
            chunk_num += 1
            logging.info(f"Processing CSV chunk {chunk_num} ({len(chunk_df)} rows)")
           
            # Add headers for first chunk
            if chunk_num == 1:
                headers = " | ".join(str(col) for col in chunk_df.columns)
                text_parts.append(f"Headers: {headers}")
           
            # Convert chunk to text
            chunk_text = dataframe_to_text(chunk_df, f"chunk_{chunk_num}")
            text_parts.append(chunk_text)
           
            # Memory management
            if chunk_num % 10 == 0:
                resources = check_system_resources()
                if resources["memory_percent"] > config.memory_threshold:
                    logging.warning(f"High memory usage processing CSV: {resources['memory_percent']:.1f}%")
                    gc.collect()
       
        result = "\n\n".join(text_parts)
        logging.info(f"Processed {chunk_num} chunks, extracted {len(result)} characters from CSV")
        return result
       
    except Exception as e:
        logging.error(f"Error processing CSV file {filename}: {e}")
        raise
 
def extract_docx_text(file, filename: str) -> str:
    """Enhanced DOCX processing"""
    logging.info(f"Processing DOCX file: {filename}")
   
    try:
        doc = Document(file)
        text_parts = []
       
        # Extract paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
           
            # Memory check for very large documents
            if i % 1000 == 0 and i > 0:
                logging.info(f"Processed {i} paragraphs")
                resources = check_system_resources()
                if resources["memory_percent"] > config.memory_threshold:
                    gc.collect()
       
        # Extract tables if any
        if doc.tables:
            logging.info(f"Processing {len(doc.tables)} tables from DOCX")
            for table_num, table in enumerate(doc.tables):
                table_text = f"[Table {table_num + 1}]"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text += f"\n{row_text}"
                text_parts.append(table_text)
       
        result = "\n\n".join(text_parts)
        logging.info(f"Extracted {len(result)} characters from DOCX")
        return result
       
    except Exception as e:
        logging.error(f"Error processing DOCX file {filename}: {e}")
        raise
 
def extract_txt_text(file, filename: str) -> str:
    """Enhanced text file processing"""
    try:
        # Handle both file objects and file paths
        if hasattr(file, 'read'):
            content = file.read()
            if isinstance(content, bytes):
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        result = content.decode(encoding)
                        logging.info(f"Successfully decoded text file with {encoding}")
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # If all encodings fail, use utf-8 with error handling
                    result = content.decode('utf-8', errors='replace')
                    logging.warning("Used UTF-8 with error replacement for text file")
            else:
                result = content
        else:
            # File path
            with open(file, 'r', encoding='utf-8') as f:
                result = f.read()
       
        logging.info(f"Extracted {len(result)} characters from text file")
        return result
       
    except Exception as e:
        logging.error(f"Error processing text file {filename}: {e}")
        raise
 
def extract_json_text(file, filename: str) -> str:
    """Enhanced JSON processing"""
    try:
        if hasattr(file, 'read'):
            content = file.read()
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            data = json.loads(content)
        else:
            with open(file, 'r') as f:
                data = json.load(f)
       
        result = json.dumps(data, indent=2, ensure_ascii=False)
        logging.info(f"Extracted {len(result)} characters from JSON file")
        return result
       
    except Exception as e:
        logging.error(f"Error processing JSON file {filename}: {e}")
        raise
 
def dataframe_to_text(df: pd.DataFrame, section_name: str = None) -> str:
    """Convert DataFrame to structured text format"""
    text_parts = []
   
    if section_name:
        text_parts.append(f"[{section_name}]")
   
    # Add data rows
    for index, row in df.iterrows():
        try:
            row_text = " | ".join(str(val) if pd.notna(val) else "N/A" for val in row.values)
            text_parts.append(f"Row {index}: {row_text}")
        except Exception as e:
            logging.warning(f"Error processing row {index}: {e}")
            continue
   
    return "\n".join(text_parts)
 
async def extract_text_async(file, filename: str) -> str:
    """
    Async wrapper for text extraction - useful for processing multiple files
    """
    loop = asyncio.get_event_loop()
   
    # Run the synchronous extraction in a thread pool
    with ThreadPoolExecutor(max_workers=1) as executor:
        result = await loop.run_in_executor(executor, extract_text, file, filename)
   
    return result
 
def get_processing_stats() -> Dict[str, Any]:
    """Get current processing statistics"""
    resources = check_system_resources()
   
    return {
        "system_resources": resources,
        "config": {
            "max_file_size_mb": config.max_file_size_mb,
            "memory_threshold": config.memory_threshold,
            "excel_chunk_size": config.excel_chunk_size,
            "csv_chunk_size": config.csv_chunk_size
        },
        "recommendations": get_performance_recommendations(resources)
    }
 
def get_performance_recommendations(resources: Dict[str, float]) -> List[str]:
    """Get performance recommendations based on current system state"""
    recommendations = []
   
    if resources["memory_percent"] > 80:
        recommendations.append("High memory usage - consider reducing chunk sizes")
   
    if resources["memory_available_gb"] < 1:
        recommendations.append("Low available memory - consider processing files individually")
   
    if resources["cpu_percent"] > 90:
        recommendations.append("High CPU usage - consider reducing concurrent operations")
   
    return recommendations
 
# Configuration functions for dynamic adjustment
def adjust_config_for_large_files():
    """Adjust configuration for processing very large files"""
    config.excel_chunk_size = 500  # Smaller chunks
    config.csv_chunk_size = 2000   # Smaller chunks
    config.memory_threshold = 70   # More conservative threshold
    logging.info("Configuration adjusted for large file processing")
 
def adjust_config_for_speed():
    """Adjust configuration for faster processing of smaller files"""
    config.excel_chunk_size = 2000  # Larger chunks
    config.csv_chunk_size = 10000   # Larger chunks
    config.memory_threshold = 85    # Less conservative threshold
    logging.info("Configuration adjusted for speed optimization")
 
def reset_config():
    """Reset configuration to defaults"""
    global config
    config = FileProcessingConfig()
    logging.info("Configuration reset to defaults")