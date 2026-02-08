import concurrent
import PyPDF2
import docx
from io import BytesIO
import logging
import os
import tempfile
import base64
from pdf2image import convert_from_path
from openai import OpenAI
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import pandas as pd
from docx import Document
import io


def estimate_text_chars(file_content: bytes, filename: str) -> int:
    file_ext = os.path.splitext(filename)[1].lower()

    try:
        if file_ext == ".pdf":
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            pages = [(p.extract_text() or "") for p in pdf_reader.pages]
            return len("\n".join(pages).strip())

        elif file_ext in [".docx", ".doc"]:
            try:
                d = docx.Document(BytesIO(file_content))
                text = "\n".join(p.text for p in d.paragraphs)
                return len(text.strip())
            except Exception:
                return len(file_content.decode("utf-8", errors="ignore"))

        elif file_ext in [".xlsx", ".xls"]:
            # read all sheets, convert to string, count chars
            xl = pd.ExcelFile(BytesIO(file_content))
            all_text = []
            for sheet in xl.sheet_names:
                df = xl.parse(sheet)
                all_text.append(df.to_csv(index=False))
            return len("\n".join(all_text))
        elif file_ext in [".csv"]:
            df = pd.read_csv(BytesIO(file_content))
            return len(df.to_csv(index=False))
        else:
            return len(file_content.decode("utf-8", errors="ignore"))

    except Exception:
        return 0

 
def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Convert DOCX bytes to PDF bytes.
    Extracts text from DOCX and writes it to a simple PDF.
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))
        
        pdf_buffer = io.BytesIO()
        pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
 
        for para in doc.paragraphs:
            if para.text.strip():
                # Handling some basic styling (very simplified)
                story.append(Paragraph(para.text, styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Handle tables in docx (basic text extraction)
        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)
            
            if data:
                t = Table(data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(t)
                story.append(Spacer(1, 12))
 
        pdf.build(story)
        pdf_buffer.seek(0)
        return pdf_buffer.read()
 
    except Exception as e:
        print(f"Error converting DOCX to PDF: {e}")
        # Fallback: Just return empty bytes or handle error appropriately
        return b""
 
def convert_excel_to_pdf(excel_bytes: bytes) -> bytes:
    """
    Convert Excel (XLSX/XLS) bytes to PDF bytes.
    Converts each sheet to a table in PDF.
    """
    try:
        # Read all sheets
        xls = pd.ExcelFile(io.BytesIO(excel_bytes))
        
        pdf_buffer = io.BytesIO()
        pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
 
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            
            # Add Sheet Title
            story.append(Paragraph(f"Sheet: {sheet_name}", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Convert DataFrame to list of lists for Table
            # Handle NaN values
            df = df.fillna('')
            data = [df.columns.values.tolist()] + df.values.tolist()
            
            # Create Table
            # Note: Large tables might break layout, this is a basic implementation
            t = Table(data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8), # Smaller font for data
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ]))
            story.append(t)
            story.append(Spacer(1, 20))
 
        pdf.build(story)
        pdf_buffer.seek(0)
        return pdf_buffer.read()
 
    except Exception as e:
        print(f"Error converting Excel to PDF: {e}")
        return b""
 
def convert_csv_to_pdf(csv_bytes: bytes) -> bytes:
    """
    Convert CSV bytes to PDF bytes.
    """
    try:
        df = pd.read_csv(io.BytesIO(csv_bytes))
        
        pdf_buffer = io.BytesIO()
        pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
 
        # Add Title
        story.append(Paragraph("CSV Data Export", styles['Heading1']))
        story.append(Spacer(1, 12))
 
        # Handle NaN values
        df = df.fillna('')
        data = [df.columns.values.tolist()] + df.values.tolist()
 
        t = Table(data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ]))
        story.append(t)
        
        pdf.build(story)
        pdf_buffer.seek(0)
        return pdf_buffer.read()
 
    except Exception as e:
        print(f"Error converting CSV to PDF: {e}")
        return b""

def encode_image(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')
    
def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text using GPT-4o Vision for maximum accuracy on tables and layouts.
    Converts PDF -> Images -> Vision API -> Markdown.
    Processes ALL pages dynamically (any PDF size) with safe batching + retries.
    """
    import os, tempfile, logging, base64, time, concurrent.futures
    from io import BytesIO
    from openai import OpenAI
    from pdf2image import convert_from_path

    file_ext = os.path.splitext(filename)[1].lower()
    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        raise ValueError("OPENAI_API_KEY not found.")

    client = OpenAI(api_key=api_key)

    try:
        # Pre-process: Convert other formats to PDF first
        if file_ext in [".docx", ".doc"]:
            print(f"🔄 Converting {filename} to PDF...")
            file_content = convert_docx_to_pdf(file_content)
            file_ext = ".pdf"
            filename = os.path.splitext(filename)[0] + ".pdf"

        elif file_ext in [".xlsx", ".xls"]:
            print(f"🔄 Converting {filename} to PDF...")
            file_content = convert_excel_to_pdf(file_content)
            file_ext = ".pdf"
            filename = os.path.splitext(filename)[0] + ".pdf"

        elif file_ext in [".csv"]:
            print(f"🔄 Converting {filename} to PDF...")
            file_content = convert_csv_to_pdf(file_content)
            file_ext = ".pdf"
            filename = os.path.splitext(filename)[0] + ".pdf"

        # Use GPT-4o Vision for PDFs
        if file_ext == ".pdf":
            print(f"👁️ Processing {filename} with GPT-4o Vision (all pages, parallel workers=5)...")

            tmp_path = None
            try:
                with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as tmp_file:
                    tmp_file.write(file_content)
                    tmp_path = tmp_file.name

                # 1) Convert PDF to images (ALL pages)
                images = convert_from_path(tmp_path)
                markdown_pages = [None] * len(images)

                print(f"📄 Converted PDF to {len(images)} images.")

                # Controls for big PDFs
                BATCH_SIZE = 10     # pages per batch (still processes ALL pages)
                MAX_WORKERS = 5     # parallel calls
                RETRIES = 3         # retry per page on transient failures
                MAX_TOKENS = 8000   # safer for dense tables

                def process_page(idx_img):
                    idx, img = idx_img

                    buf = BytesIO()
                    img.save(buf, format="JPEG")
                    base64_image = base64.b64encode(buf.getvalue()).decode("utf-8")
                    buf.close()

                    for attempt in range(RETRIES):
                        try:
                            response = client.chat.completions.create(
                                model="gpt-4o",
                                messages=[
                                    {
                                        "role": "user",
                                        "content": [
                                            {
                                                "type": "text",
                                                "text": (
                                                    "Transcribe this page into Markdown. "
                                                    "Preserve all tables, headers, and layout structures exactly. "
                                                    "If there is a table, output it as a Markdown table. "
                                                    "Do not summarize."
                                                ),
                                            },
                                            {
                                                "type": "image_url",
                                                "image_url": {
                                                    "url": f"data:image/jpeg;base64,{base64_image}",
                                                    "detail": "high"
                                                },
                                            },
                                        ],
                                    }
                                ],
                                max_tokens=MAX_TOKENS,
                                timeout=300
                            )

                            choice = response.choices[0]
                            if choice.finish_reason == "length":
                                logging.warning(f"⚠️ Page {idx+1} may be truncated (token limit hit).")

                            page_text = choice.message.content
                            print(f"✅ Processed page {idx+1}/{len(images)}")
                            return idx, page_text

                        except Exception as e:
                            if attempt == RETRIES - 1:
                                raise
                            wait = 2 ** attempt
                            logging.warning(
                                f"Retrying page {idx+1} after error: {e}. Waiting {wait}s"
                            )
                            time.sleep(wait)

                # 2) Run in batches to avoid rate limits / memory spikes
                for start in range(0, len(images), BATCH_SIZE):
                    batch_imgs = images[start:start + BATCH_SIZE]
                    batch = list(enumerate(batch_imgs, start=start))

                    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                        futures = [executor.submit(process_page, item) for item in batch]
                        for future in concurrent.futures.as_completed(futures):
                            idx, text = future.result()
                            markdown_pages[idx] = text

                full_text = "\n\n".join(markdown_pages)

                # Clean log (don’t dump full KB)
                logging.info(
                    f"✅ Completed Vision extraction for {full_text} "
                    f"(pages={len(images)}, chars={len(full_text)})"
                )

                return full_text

            except Exception as e:
                logging.error(f"Vision extraction failed for {filename}: {e}")
                raise

            finally:
                if tmp_path and os.path.exists(tmp_path):
                    os.remove(tmp_path)

        else:
            # Fallback for non-PDF types
            return file_content.decode("utf-8", errors="ignore")

    except Exception as e:
        print(f"Error extracting text from {filename}: {str(e)}")
        return ""
